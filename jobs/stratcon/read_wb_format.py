from docx import Document
import unicodedata as ud
import glob
import re
from base.utils.SQLConnection import MSSqlConnection
from base.utils.Logger import Logger
# from docx.text import run
# from functools import reduce

log = Logger("Strategy Consulting").getlog()

_PATH_SEP = "\\"
_HOME_DRIVE = "C:\\Users\\gunjan.kumar"
_TEAM_NAME = "StrategicConsulting"
_PROJ_DIR = _HOME_DRIVE + _PATH_SEP + "Documents\\JLL\\Projects"
_TEAM_DIR = _PROJ_DIR + _PATH_SEP + _TEAM_NAME
_DOC_PATH = _TEAM_DIR + _PATH_SEP + "ProjectSheetsEvaluations"

_COL_FILE_NAME = "FileNames"
_COL_TABLE_KEY = "TableKey"
_COL_ADDRESS = "ClientAddress"
_COL_ASSIGNMENT_DURATION = "AssignmentDuration"
_COL_ASSIGNMENT_NAME = "AssignmentName"
_COL_ASSOCIATED_CONSULT = "AssociatedConsultant"
_COL_CLIENT_NAME = "ClientName"
_COL_CLIENT_TYPE = "ClientType"
_COL_COUNTRY_NAME = "CountryName"
_COL_END_DATE = "EndDate"
_COL_JLL_RECOMMENDATION = "JLLRecommendation"
_COL_NO_OF_STAFF = "NoOfStaff"
_COL_SCOPE_OF_WORK = "ScopeOfWork"
_COL_SERVICE_VALUE = "ServiceValue"
_COL_SITE_LOCATION = "LocationOfSite"
_COL_STAFF_MONTH = "StaffMonth"
_COL_START_DATE = "StartDate"
_COL_STUDY_TYPE = "StudyType"
_COL_TEAM_NAME = "TeamName"
_COL_PROJECT_VALUE = "ProjectValue"
_COL_ASSOCIATED_FIRM = "AssociatedFirmName"

_ALL_COLUMNS = [
    _COL_FILE_NAME,
    _COL_TABLE_KEY,
    _COL_ADDRESS,
    _COL_ASSIGNMENT_DURATION,
    _COL_ASSIGNMENT_NAME,
    _COL_ASSOCIATED_CONSULT,
    _COL_CLIENT_NAME,
    _COL_CLIENT_TYPE,
    _COL_COUNTRY_NAME,
    _COL_END_DATE,
    _COL_JLL_RECOMMENDATION,
    _COL_NO_OF_STAFF,
    _COL_SCOPE_OF_WORK,
    _COL_SERVICE_VALUE,
    _COL_SITE_LOCATION,
    _COL_STAFF_MONTH,
    _COL_START_DATE,
    _COL_STUDY_TYPE,
    _COL_TEAM_NAME,
    _COL_PROJECT_VALUE,
    _COL_ASSOCIATED_FIRM
]

_FIELD_NAME = "Field"
_FIELD_VALUE = "Value"

_FILE_FORMAT_MAP = {
    "docx": {
        "World Bank": {
            "FILENAME": "WB",
            "TYPEOFDATA": "TABLE",
            "FIELDMAP": {
                _COL_ADDRESS: ["Address and Contact Person (Client)",
                               "Address"],
                _COL_ASSIGNMENT_DURATION: ["Duration of Assignment"],
                _COL_ASSIGNMENT_NAME: ["Project Name",
                                       "Assignment Name"],
                _COL_ASSOCIATED_CONSULT: ["No of Person",
                                          "No. of Person-Months of Professional"
                                          " Staff Provided by Associated "
                                          "Firm(s)",
                                          "No. of months of key professional"
                                          " staff, provided by Associated"
                                          " Consultants"],
                _COL_ASSOCIATED_FIRM: ["Name of the Associated Firm(s), if any",
                                       "Name of associated firm(s) if any",
                                       "Name of Associated Consultants, if any"
                                       ],
                _COL_CLIENT_NAME: ["Name of Client",
                                   "Name of the Client"],
                _COL_CLIENT_TYPE: [""],
                _COL_COUNTRY_NAME: ["Country"],
                _COL_END_DATE: ["Completion Date"],
                _COL_JLL_RECOMMENDATION: [""],
                _COL_NO_OF_STAFF: ["Professional Staff Provided by your Firm",
                                   "No. of Staff",
                                   "Professional Staff Provided by our Company"
                                   ],
                _COL_PROJECT_VALUE: ["Approximate Value of Project",
                                     "Approx. Project Cost"],
                _COL_SCOPE_OF_WORK: ["Scope of work",
                                     "Detailed Description of Services "
                                     "Provided",
                                     "Description of Actual Service provided by"
                                     " your Staff within the assignment",
                                     "Description of Actual Services Provided"
                                     " by your Staff"],
                _COL_SERVICE_VALUE: ["Approximate Value of Services",
                                     "Value of Contract",
                                     "Value of Contract (in USD)",
                                     "Approx. Value of Services"],
                _COL_SITE_LOCATION: ["Project Location within the Country",
                                     "Project Locations within Country",
                                     "Name of City",
                                     "Location within Country"],
                _COL_STAFF_MONTH: ["No of Staff Month",
                                   "No of Person Month",
                                   "Total No of Staff Month",
                                   "No of Person Months",
                                   "No. of Staff-Months"],
                _COL_START_DATE: ["Start Date"],
                _COL_TEAM_NAME: ["Name of Senior Staff (Project Director/"
                                 "Coordinator, Team Leader) involved and "
                                 "functions performed"],
                _COL_STUDY_TYPE: ["Narrative Description of Project",
                                  "Detailed Narrative Description of Project",
                                  "Narrative Description"]
            }
        }
    }
}


_DELETE_WORD = ["the", "-", "\."]
_CHECK_LENGTH = {}


class ReadWordDocument:
    def __init__(self):
        self.filename = None
        self.doc = None
        self.tables = None

    def readfile(self, filename):
        self.filename = filename
        self.doc = Document(self.filename)
        self.tables = self.doc.tables

    def _read_table(self, celloverwrite=False, cellappend=False,
                    cellappdelim=";", paraappend=True, paraappdelim=" ",
                    paraoverwrite=False):
        __ddetail = {}
        __vtable = 1
        for __ttable in self.tables:
            __dtemp = {}
            for __trow in __ttable.rows:
                for __tcell in __trow.cells:
                    __vpara = ""
                    for __tpara in __tcell.paragraphs:
                        if paraoverwrite:
                            __vpara = __tpara.text
                        elif paraappend:
                            __vpara += __tpara.text + paraappdelim
                    __vpara.strip()
                    __vkey = __vpara.split(":")[0].strip()
                    __vvalue = ud.normalize("NFKD", " ".join(
                        __vpara.split(":")[1:])
                    )
                    __vkey = __vkey.strip()
                    __vvalue = __vvalue.strip()
                    if __vkey in __dtemp:
                        if celloverwrite:
                            __dtemp[__vkey] = __vvalue
                        elif cellappend:
                            __dtemp[__vkey] = __dtemp[__vkey] +\
                                              cellappdelim +\
                                              __vvalue
                    else:
                        __dtemp[__vkey] = __vvalue
            __ddetail[__vtable] = __dtemp
            __vtable += 1
        return __ddetail

    def get_data(self, tabledata=True, **kwargs):
        __dtable = {}
        if tabledata:
            __dtable = self._read_table(**kwargs)
        return __dtable


def readdocx(basedir=_DOC_PATH, filename=None, pathsep=_PATH_SEP,
             wordformat=None):
    __odoc = ReadWordDocument()
    __dfile = {}
    for key, value in wordformat.items():
        __lfile = []
        __btable = False
        if value["TYPEOFDATA"] == "TABLE":
            __btable = True

        def _read_doc_data(files):
            __dtemp = {}
            __odoc.readfile(files)
            __vfile = files[files.rfind(pathsep)+1:]
            __dtemp[__vfile] = __odoc.get_data(tabledata=__btable)
            return __dtemp
        if filename:
            __lfile.append(_read_doc_data(filename))
        else:
            for __vdocx in glob.glob(basedir + pathsep + value["FILENAME"] +
                                     "*.docx"):
                log.info("..Processing File: {}".format(__vdocx))
                __lfile.append(_read_doc_data(__vdocx))
        __dfile[key] = __lfile
    return __dfile


def clean_string(string):
    __vstr = re.sub(r'[^\x00-\x7f]', r'', string)
    __vstr = re.sub(r'\s+', r' ', __vstr)
    __vstr = __vstr.strip()
    return __vstr


def get_column_name(colname, fieldmap=None):
    # __ore = re.compile('\s+')
    # __string = reduce(lambda x, y: x + "|" + y, _DELETE_WORD)
    # __vcol = re.sub(r'(' + __string + ')', r'', colname, flags=re.IGNORECASE)
    # __vcol = re.sub(__ore, ' ', __vcol)
    __retcol = ""
    for key, value in fieldmap.items():
        __vmatch = False
        for __tdetail in value:
            if __tdetail == "":
                continue
            # __vtemp = re.sub(r'(' + __string + ')', r'', __tdetail,
            #                  flags=re.IGNORECASE)
            # __vtemp = re.sub(__ore, ' ', __vtemp)
            # if re.match(r'' + __vtemp + '', __vcol, flags=re.IGNORECASE):
            #     __vmatch = True
            if colname.upper() == __tdetail.upper():
                __vmatch = True
                break
        if __vmatch:
            __retcol = key
            break
    return __retcol


def converttime(curtime, curtype, tgttype="Month"):
    __vreturn = "NULL"
    __vxform = 1
    if tgttype:
        if tgttype.upper() == "MONTH":
            if curtype.upper() in ("YEAR", "YEARS"):
                __vxform = 12
            elif curtype.upper() in ("DAY", "DAYS"):
                __vxform = (1/30)
    if curtime and __vxform:
        __vreturn = curtime * __vxform
    return __vreturn


def splitstring(string, keyname):
    def _split_time(pstring):
        __oretime = re.compile(r'(year|month|day)', flags=re.IGNORECASE)
        __orenumber = re.compile(r'\d+')
        __ldata = map(lambda x: x.strip(), pstring.split(" "))
        __vnum = None
        __vtype = None
        for __tdata in __ldata:
            if __orenumber.match(__tdata):
                if not __vnum:
                    __vnum = __tdata
            else:
                __orematch = __oretime.match(__tdata)
                if __orematch:
                    __vtype = __orematch.group(1)
        __vvalue = converttime(__vnum, __vtype)
        # if __vmult and __vnum:
        #     __vvalue = __vnum * __vmult
        return __vvalue

    def _split_manpower(pstring):
        __vstaffnum = "NULL"
        __vnummonth = "NULL"
        __orestr = re.compile(r"^([^\d]+staff[^\d]+)(\d*)"
                              r"([^\d]+(month|year|day)[^\d]*)(\d*)",
                              flags=re.IGNORECASE)
        __orestr2 = re.compile(r"^([^\d]+staff[^\d]+)(\d+)",
                               flags=re.IGNORECASE)
        __orematch = __orestr.match(pstring)
        if __orematch:
            __vstaffnum = __orematch.group(2)
            if not __vstaffnum or __vstaffnum == "":
                __vstaffnum = "NULL"
            try:
                __vtype = __orematch.group(4)
                __vtime = __orematch.group(5)
                if __vtime and __vtype:
                    __vnummonth = converttime(__vtime, __vtype)
                else:
                    __vnummonth = "NULL"
            except Exception as err:
                __vnummonth = "NULL"
        else:
            __orematch = __orestr2.match(pstring)
            if __orematch:
                __vstaffnum = __orematch.group(2)
                if not __vstaffnum or __vstaffnum == "":
                    __vstaffnum = "NULL"
            else:
                if pstring:
                    __vstaffnum = pstring
                    __vnummonth = "NULL"
        return [__vstaffnum, __vnummonth]

    __lreturn = None
    if keyname == _COL_NO_OF_STAFF:
        __lreturn = _split_manpower(string)
    elif keyname == _COL_ASSIGNMENT_DURATION:
        __lreturn = _split_time(string)
    return __lreturn


def column_assignment(filename, tablekey, record, colmap=None):
    __dcolumn = {
        _COL_FILE_NAME: filename,
        _COL_TABLE_KEY: tablekey
    }
    for key, value in record.items():
        __ovalue = None
        __tkey = clean_string(key)
        if __tkey == "":
            continue
        __tkey = get_column_name(__tkey, fieldmap=colmap["FIELDMAP"])
        if __tkey == "":
            continue
        __tvalue = clean_string(value)
        if __tvalue == "" or __tvalue.upper() in ("NA", "NIL", "NONE"):
            __tvalue = "NULL"
        if __tkey in (_COL_NO_OF_STAFF, _COL_ASSIGNMENT_DURATION):
            __ovalue = splitstring(__tvalue, __tkey)
            if __tkey == _COL_ASSIGNMENT_DURATION:
                __dcolumn[_COL_ASSIGNMENT_DURATION] = __ovalue
            else:
                __dcolumn[_COL_NO_OF_STAFF] = __ovalue[0]
                if _COL_STAFF_MONTH not in __dcolumn:
                    __dcolumn[_COL_STAFF_MONTH] = __ovalue[1]
        else:
            if __tkey == _COL_STAFF_MONTH:
                if __tkey in __dcolumn:
                    if __tvalue == "NULL" and __dcolumn[__tkey] != "NULL":
                        continue

            __dcolumn[__tkey] = __tvalue
        if __tkey in _CHECK_LENGTH:
            if len(__tvalue) > _CHECK_LENGTH[__tkey]:
                _CHECK_LENGTH[__tkey] = len(__tvalue)
        else:
            _CHECK_LENGTH[__tkey] = len(__tvalue)
    for __lcol in _ALL_COLUMNS:
        if __lcol not in __dcolumn:
            __dcolumn[__lcol] = "NULL"
    return __dcolumn


def make_insert(filename, tablekey, record, tablename=None, colmap=None):
    __vcolstr = "FileNames, TableKey, "
    __vvalstr = "'" + filename + "', '" + str(tablekey) + "', "
    for key, value in record.items():
        __tkey = clean_string(key)
        if __tkey == "":
            continue
        __tkey = get_column_name(__tkey, fieldmap=colmap["FIELDMAP"])
        if __tkey == "":
            continue
        __vcolstr += __tkey + ", "
        __tvalue = clean_string(value)
        if __tvalue == "":
            __vvalstr += "NULL, "
        else:
            __vvalstr += "'" + __tvalue + "', "
        if __tkey in _CHECK_LENGTH:
            if len(__tvalue) > _CHECK_LENGTH[__tkey]:
                _CHECK_LENGTH[__tkey] = len(__tvalue)
        else:
            _CHECK_LENGTH[__tkey] = len(__tvalue)
    __vcolstr = __vcolstr[:__vcolstr.rfind(",")]
    __vvalstr = __vvalstr[:__vvalstr.rfind(",")]
    __vinsert = "INSERT INTO {0} ({1}) VALUES ({2})".format(
        tablename,
        __vcolstr,
        __vvalstr
    )
    return __vinsert


def prepare_loading_data():
    __dmap = _FILE_FORMAT_MAP["docx"]
    __vdata = readdocx(_DOC_PATH, wordformat=__dmap)
    # __vdata = readdocx(filename=_DOC_PATH + _PATH_SEP +
    #                    "WB Format_Chennai.docx",
    #                    wordformat=__dmap)
    __linsert = []
    for __vformat, __ldetail in __vdata.items():
        for __tdetail in __ldetail:
            for __filekey, __filevalue in __tdetail.items():
                for __tabkey, __tabvalue in __filevalue.items():
                    __linsert.append(
                        column_assignment(
                            __filekey,
                            __tabkey,
                            __tabvalue,
                            colmap=__dmap[__vformat]
                        )
                    )
    return __linsert


def test_connection():
    __vserver = "uslilvmdwstg1.delphiprd.am.joneslanglasalle.com"
    __vdatabase = "GDIM_india"
    __oconn = MSSqlConnection(__vserver, database=__vdatabase)
    __oconn.connect()
    log.info("Connection Established and Successful")
    __oconn.disconnect()
    log.info("Connection Closed")


def load_data():
    __vtable = "GDIM_India.IN_LOCAL_RAW.StratConEvaluate"
    __linsert = prepare_loading_data()
    __lallcols = _ALL_COLUMNS
    __lallcols.sort()
    __vquery = "INSERT INTO {0} ({1}) VALUES ({2})"
    __vquery = __vquery.format(
        __vtable,
        ", ".join(str(x) for x in __lallcols),
        ", ".join("?" * len(__lallcols))
    )
    # print(_CHECK_LENGTH)
    __vserver = "uslilvmdwstg1.delphiprd.am.joneslanglasalle.com"
    __vdatabase = "GDIM_india"
    __oconn = MSSqlConnection(__vserver, database=__vdatabase)
    __oconn.connect()
    try:
        __vline = 1
        for __tinsert in __linsert:
            __ldata = []
            log.info("..Processing Line {}".format(str(__vline)))
            for __lcol in __lallcols:
                __ldata.append(__tinsert[__lcol])
            __oconn.cursor.execute(__vquery, __ldata)
            __vline += 1
        __oconn.conn.commit()
    except Exception as err:
        log.info("Insert Failed: Exiting")
        log.error(err)
        pass
    __oconn.disconnect()


load_data()
# test_connection()
# def insert_data(inslist):
# for __tinsert in linsert:
#     log.info(__tinsert)
